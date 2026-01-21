import os
import logging
import uuid
import glob
from typing import List, Dict, Any, Optional
from datetime import datetime
import time
from tqdm import tqdm

import chromadb
from chromadb.config import Settings as ChromaSettings
from chromadb.utils import embedding_functions
from sentence_transformers import SentenceTransformer

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import pypdf
from docx import Document as DocxDocument
import anthropic
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

from backend.config import settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles document loading, chunking, embedding, and indexing into ChromaDB.
    """
    
    def __init__(self, persistence_path: str = "data/chroma_db"):
        """
        Initialize the DocumentProcessor.
        
        Args:
            persistence_path (str): Path to persist ChromaDB data.
        """
        self.persistence_path = persistence_path
        
        # Initialize ChromaDB Client
        # Using persistent client
        logger.info(f"Initializing ChromaDB Client at {self.persistence_path}...")
        try:
            self.chroma_client = chromadb.PersistentClient(path=self.persistence_path)
        except Exception as e:
            logger.error(f"Failed to initialize ChromaDB: {e}")
            raise

        # Initialize Embedding Model
        # using the requested model: sentence-transformers/all-MiniLM-L6-v2
        model_name = "sentence-transformers/all-MiniLM-L6-v2"
        logger.info(f"Loading embedding model: {model_name}...")
        try:
            self.embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=model_name)
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            raise
            
        # Get or Create Collection
        self.collection_name = "company_docs"
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            embedding_function=self.embedding_fn
        )
        logger.info(f"Collection '{self.collection_name}' ready. Count: {self.collection.count()}")

    def load_documents(self, folder_path: str) -> List[Document]:
        """
        Load documents from the specified folder path recursively.
        Supports: .pdf, .docx, .txt, .md
        
        Args:
            folder_path (str): Path to the documents folder.
            
        Returns:
            List[Document]: List of LangChain Document objects.
        """
        documents = []
        folder_path = os.path.abspath(folder_path)
        
        if not os.path.exists(folder_path):
            logger.warning(f"Folder path does not exist: {folder_path}")
            return []

        # Supported extensions
        extensions = ['*.pdf', '*.docx', '*.txt', '*.md']
        files = []
        for ext in extensions:
            # Recursive search for each extension
            files.extend(glob.glob(os.path.join(folder_path, '**', ext), recursive=True))
            
        logger.info(f"Found {len(files)} documents in {folder_path}")
        
        for file_path in tqdm(files, desc="Loading documents"):
            try:
                ext = os.path.splitext(file_path)[1].lower()
                content = ""
                metadata = {
                    "source": os.path.basename(file_path),
                    "path": file_path,
                    "size": os.path.getsize(file_path),
                    "type": ext
                }
                
                if ext == '.pdf':
                    reader = pypdf.PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
                            
                elif ext == '.docx':
                    doc = DocxDocument(file_path)
                    content = "\n".join([para.text for para in doc.paragraphs])
                    
                elif ext in ['.txt', '.md']:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                
                if content.strip():
                    documents.append(Document(page_content=content, metadata=metadata))
                    
            except Exception as e:
                logger.error(f"Error loading file {file_path}: {e}")
                continue
                
        logger.info(f"Successfully loaded {len(documents)} documents.")
        return documents

    def chunk_documents(self, documents: List[Document], chunk_size: int = 500, overlap: int = 50) -> List[Document]:
        """
        Split documents into chunks.
        
        Args:
            documents (List[Document]): List of documents to chunk.
            chunk_size (int): Max token/character size for chunks.
            overlap (int): Overlap between chunks.
            
        Returns:
            List[Document]: List of chunked documents.
        """
        if not documents:
            return []
            
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len, # Can use tiktoken if precise token counting is needed, but len is fast approx
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_documents(documents)
        logger.info(f"Split {len(documents)} documents into {len(chunks)} chunks.")
        return chunks

    def index_documents(self, chunks: List[Document]) -> int:
        """
        Index chunks into ChromaDB.
        
        Args:
            chunks (List[Document]): List of document chunks.
            
        Returns:
            int: Number of chunks indexed.
        """
        if not chunks:
            logger.warning("No chunks to index.")
            return 0
            
        logger.info(f"Indexing {len(chunks)} chunks into ChromaDB...")
        
        ids = []
        documents_content = []
        metadatas = []
        
        for chunk in tqdm(chunks, desc="Preparing chunks"):
            # Generate a unique ID for each chunk
            chunk_id = str(uuid.uuid4())
            ids.append(chunk_id)
            
            documents_content.append(chunk.page_content)
            
            # Enrich metadata
            meta = chunk.metadata.copy()
            meta["chunk_id"] = chunk_id
            meta["timestamp"] = datetime.now().isoformat()
            metadatas.append(meta)
            
        # Add to collection in batches (Chroma handles batching reasonably well, but explicit batching is safer for huge datasets)
        batch_size = 100
        for i in tqdm(range(0, len(ids), batch_size), desc="Indexing batches"):
            end_idx = i + batch_size
            try:
                self.collection.add(
                    documents=documents_content[i:end_idx],
                    metadatas=metadatas[i:end_idx],
                    ids=ids[i:end_idx]
                )
            except Exception as e:
                logger.error(f"Error indexing batch {i}-{end_idx}: {e}")
                
        count = self.collection.count()
        logger.info(f"Indexing complete. Total documents in collection: {count}")
        return len(ids)

    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Search for relevant documents in ChromaDB.
        
        Args:
            query (str): The query string.
            top_k (int): Number of top results to return.
            
        Returns:
            List[Dict[str, Any]]: List of results with content and metadata.
        """
        if not query:
            return []
            
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )
            
            # Formatting results
            formatted_results = []
            if results and results['documents']:
                for i in range(len(results['documents'][0])):
                    formatted_results.append({
                        "content": results['documents'][0][i],
                        "metadata": results['metadatas'][0][i],
                        "similarity_score": 1.0 - (results['distances'][0][i] if results['distances'] else 0), # Approx convert distance to similarity if using cosine/l2
                        "source": results['metadatas'][0][i].get("source", "Unknown")
                    })
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error during search: {e}")
            return []

class ClaudeRAG:
    """
    RAG engine using Anthropic's Claude to answer questions based on retrieved documentation.
    """

    def __init__(self, persistence_path: str = "data/chroma_db"):
        """
        Initialize the ClaudeRAG engine.

        Args:
            persistence_path (str): Path to persist ChromaDB data.
        """
        self.doc_processor = DocumentProcessor(persistence_path=persistence_path)
        
        # Initialize Anthropic Client
        api_key = settings.ANTHROPIC_API_KEY
        if not api_key:
            logger.warning("ANTHROPIC_API_KEY not found in settings. API calls will fail.")
        
        self.client = anthropic.Anthropic(api_key=api_key)
        
        # Pricing configuration (Cost per million tokens)
        self.input_price = settings.CLAUDE_INPUT_PRICE_PER_MILLION
        self.output_price = settings.CLAUDE_OUTPUT_PRICE_PER_MILLION
        
        # Model configuration
        # Using the specific model requested by user
        self.model = "claude-sonnet-4-20250514" 

    def _calculate_cost(self, input_tokens: int, output_tokens: int) -> float:
        """
        Calculate cost in USD based on token usage.
        """
        input_cost = (input_tokens / 1_000_000) * self.input_price
        output_cost = (output_tokens / 1_000_000) * self.output_price
        return round(input_cost + output_cost, 6)

    @retry(
        retry=retry_if_exception_type((anthropic.APIConnectionError, anthropic.RateLimitError, anthropic.InternalServerError)),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def _call_claude_api(self, messages: List[Dict[str, str]], system_prompt: str) -> Any:
        """
        Internal method to call Claude API with retry logic.
        """
        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=2048,
                system=system_prompt,
                messages=messages
            )
            return response
        except anthropic.BadRequestError as e:
            # Do not retry on Bad Request (likely invalid request/model)
            logger.error(f"Bad Request to Anthropic API: {e}")
            raise e
        except Exception as e:
            logger.error(f"Error calling Anthropic API: {e}")
            raise e

    def ask(self, query: str, conversation_history: List[Dict[str, str]] = None) -> Dict[str, Any]:
        """
        Ask a question to the RAG system.

        Args:
            query (str): User's question.
            conversation_history (List[Dict]): Previous messages [{"role": "user", "content": "..."}, ...]

        Returns:
            Dict: Answer, sources, usage stats, and cost.
        """
        start_time = time.time()
        
        # 1. Retrieve context
        logger.info(f"Searching context for query: {query}")
        context_chunks = self.doc_processor.search(query, top_k=5)
        
        context_text = ""
        sources = set()
        for chunk in context_chunks:
            source = chunk.get('source', 'Unknown')
            sources.add(source)
            context_text += f"---\nSource: {source}\nContent: {chunk.get('content')}\n\n"
        
        # 2. Construct Prompt
        system_prompt = f"""Eres un asistente comercial experto. Responde basándote ÚNICAMENTE 
en la documentación proporcionada.

CONTEXTO DE DOCUMENTACIÓN:
{context_text}

INSTRUCCIONES:
- Responde de forma clara y profesional
- Cita las fuentes cuando uses información específica
- Si no encuentras la respuesta en la documentación, dilo claramente
- No inventes información
"""
        
        # Prepare messages
        messages = []
        if conversation_history:
            messages.extend(conversation_history)
        
        messages.append({"role": "user", "content": query})

        # 3. Call API
        try:
            logger.info("Sending request to Claude API...")
            response = self._call_claude_api(messages, system_prompt)
            
            # 4. Process Response
            answer_text = response.content[0].text
            input_tokens = response.usage.input_tokens
            output_tokens = response.usage.output_tokens
            
            cost = self._calculate_cost(input_tokens, output_tokens)
            
            logger.info(f"Answer received. Tokens: In={input_tokens}, Out={output_tokens}. Cost: ${cost}")
            
            return {
                "answer": answer_text,
                "sources": list(sources),
                "tokens_input": input_tokens,
                "tokens_output": output_tokens,
                "cost_usd": cost,
                "context_used": context_chunks,
                "elapsed_time": round(time.time() - start_time, 2)
            }
            
        except Exception as e:
            logger.error(f"RAG execution failed: {e}")
            return {
                "answer": "Lo siento, hubo un error al procesar tu solicitud. Por favor intenta más tarde.",
                "error": str(e),
                "sources": [],
                "tokens_input": 0,
                "tokens_output": 0,
                "cost_usd": 0.0,
                "context_used": []
            }

    def reindex_all(self, folder_path: str = "data/documents/") -> Dict[str, Any]:
        """
        Clear the database and re-index all documents.

        Args:
            folder_path (str): Path to documents folder.

        Returns:
            Dict: Statistics of re-indexing.
        """
        start_time = time.time()
        logger.info("Starting full re-indexing...")
        
        try:
            # 1. Reset Collection (Delete and Recreate)
            try:
                self.doc_processor.chroma_client.delete_collection(self.doc_processor.collection_name)
                logger.info("Collection deleted.")
            except Exception as e:
                logger.warning(f"Collection deletion failed (might not exist): {e}")

            # Re-initialize collection
            self.doc_processor.collection = self.doc_processor.chroma_client.get_or_create_collection(
                name=self.doc_processor.collection_name,
                embedding_function=self.doc_processor.embedding_fn
            )
            
            # 2. Load and Index
            documents = self.doc_processor.load_documents(folder_path)
            chunks = self.doc_processor.chunk_documents(documents)
            count = self.doc_processor.index_documents(chunks)
            
            elapsed = round(time.time() - start_time, 2)
            stats = {
                "status": "success",
                "documents_found": len(documents),
                "chunks_indexed": count,
                "elapsed_time_seconds": elapsed
            }
            logger.info(f"Re-indexing complete: {stats}")
            return stats
            
        except Exception as e:
            logger.error(f"Re-indexing failed: {e}")
            return {
                "status": "error",
                "error": str(e),
                "elapsed_time_seconds": round(time.time() - start_time, 2)
            }
